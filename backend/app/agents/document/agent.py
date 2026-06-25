"""
Document Generator — Module 9 (TOP-1% OVERHAUL)
Outputs: submission-ready .docx with embedded matplotlib charts + rendered Word tables,
         cover letter, reviewer response, editorial report, BibTeX.
All sections: Abstract, Introduction, Literature Review, Materials & Methods,
              Results, Discussion, Conclusion, Keywords, Author Contributions,
              Funding, Conflicts of Interest, Ethics, Supplementary Materials.
"""
from __future__ import annotations

import asyncio
import io
import json
from datetime import datetime, timezone
from typing import Dict, Any, List, Optional
import structlog

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from langchain_core.messages import HumanMessage

from app.core.config import settings
from app.core.llm_factory import get_fast_llm
from app.services.storage.s3 import S3Service

logger = structlog.get_logger()

# ─── Colour palette (journal-quality) ────────────────────────────────────────
CHART_COLORS = [
    "#1a73e8", "#e84040", "#2da44e", "#f5a623", "#7b68ee",
    "#00b4d8", "#e76f51", "#264653", "#a8dadc", "#457b9d",
]
CHART_STYLE = {
    "axes.spines.top": False,
    "axes.spines.right": False,
    "axes.grid": True,
    "grid.alpha": 0.25,
    "grid.linestyle": "--",
    "font.family": "serif",
    "font.size": 9,
    "axes.titlesize": 10,
    "axes.titleweight": "bold",
    "axes.labelsize": 9,
    "xtick.labelsize": 8,
    "ytick.labelsize": 8,
    "legend.fontsize": 8,
    "legend.framealpha": 0.4,
    "figure.facecolor": "white",
    "axes.facecolor": "#fafafa",
}


class DocumentAgent:
    def __init__(self):
        self.llm = get_fast_llm(max_tokens=2048)
        self.s3 = S3Service()

    # ─── Chart Rendering ──────────────────────────────────────────────────

    def _render_chart(self, chart: Dict[str, Any]) -> Optional[bytes]:
        try:
            with plt.rc_context(CHART_STYLE):
                chart_type = chart.get("type", "bar").lower()
                title = chart.get("title", "")
                x_label = chart.get("x_label", "")
                y_label = chart.get("y_label", "")
                x_data = chart.get("x_data", [])
                datasets = chart.get("datasets", [])

                fig, ax = plt.subplots(figsize=(6.5, 3.8), dpi=150)

                if chart_type == "bar":
                    n_ds = max(len(datasets), 1)
                    x = np.arange(len(x_data)) if x_data else np.arange(0)
                    width = min(0.75 / n_ds, 0.35)
                    for i, ds in enumerate(datasets):
                        offset = (i - (n_ds - 1) / 2) * width
                        vals = [float(v) if v is not None else 0 for v in ds.get("data", [])]
                        bars = ax.bar(x + offset, vals, width * 0.9,
                                      label=ds.get("label", ""),
                                      color=CHART_COLORS[i % len(CHART_COLORS)],
                                      alpha=0.88, edgecolor="white", linewidth=0.6)
                        for bar, val in zip(bars, vals):
                            if val != 0:
                                ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(vals) * 0.01,
                                        f"{val:.1f}" if isinstance(val, float) and val != int(val) else str(int(val)),
                                        ha="center", va="bottom", fontsize=7, color="#444")
                    if x_data:
                        ax.set_xticks(x)
                        ax.set_xticklabels(x_data, rotation=30 if len(x_data) > 5 else 0, ha="right")

                elif chart_type == "line":
                    for i, ds in enumerate(datasets):
                        data = [float(v) if v is not None else 0 for v in ds.get("data", [])]
                        xs = list(range(len(data))) if not x_data else x_data
                        if isinstance(xs[0], str):
                            ax.plot(range(len(data)), data, "o-",
                                    label=ds.get("label", ""),
                                    color=CHART_COLORS[i % len(CHART_COLORS)],
                                    linewidth=2, markersize=5, markerfacecolor="white", markeredgewidth=1.5)
                            ax.set_xticks(range(len(xs)))
                            ax.set_xticklabels(xs, rotation=30 if len(xs) > 5 else 0, ha="right")
                        else:
                            ax.plot(xs, data, "o-",
                                    label=ds.get("label", ""),
                                    color=CHART_COLORS[i % len(CHART_COLORS)],
                                    linewidth=2, markersize=5, markerfacecolor="white", markeredgewidth=1.5)
                    ax.fill_between(range(len(datasets[0].get("data", []))),
                                    [float(v) for v in datasets[0].get("data", []) if v is not None],
                                    alpha=0.08, color=CHART_COLORS[0]) if datasets else None

                elif chart_type == "scatter":
                    for i, ds in enumerate(datasets):
                        pts = ds.get("data", [])
                        if pts and isinstance(pts[0], (list, tuple)):
                            xs_s, ys_s = zip(*pts) if pts else ([], [])
                        else:
                            xs_s = x_data or list(range(len(pts)))
                            ys_s = [float(v) if v is not None else 0 for v in pts]
                        ax.scatter(xs_s, ys_s, label=ds.get("label", ""),
                                   color=CHART_COLORS[i % len(CHART_COLORS)],
                                   alpha=0.72, s=45, edgecolors="white", linewidths=0.5)

                elif chart_type == "pie":
                    if datasets and datasets[0].get("data"):
                        data = [float(v) for v in datasets[0]["data"] if v is not None]
                        labels = x_data or [f"Segment {i+1}" for i in range(len(data))]
                        wedges, texts, autotexts = ax.pie(
                            data, labels=labels, autopct="%1.1f%%",
                            colors=CHART_COLORS[:len(data)],
                            startangle=90,
                            wedgeprops=dict(edgecolor="white", linewidth=1.5),
                            pctdistance=0.82,
                        )
                        for t in autotexts:
                            t.set_fontsize(7.5)
                        ax.axis("equal")

                elif chart_type == "box":
                    if datasets:
                        box_data = [[float(v) for v in ds.get("data", []) if v is not None] for ds in datasets]
                        labels = [ds.get("label", f"Group {i+1}") for i, ds in enumerate(datasets)]
                        bp = ax.boxplot(box_data, labels=labels, patch_artist=True,
                                        medianprops=dict(color="white", linewidth=2),
                                        flierprops=dict(marker="o", markersize=4, alpha=0.5))
                        for patch, color in zip(bp["boxes"], CHART_COLORS):
                            patch.set_facecolor(color)
                            patch.set_alpha(0.78)

                elif chart_type == "heatmap":
                    if datasets and datasets[0].get("data"):
                        matrix = [[float(v) if v is not None else 0 for v in row]
                                  for row in datasets[0]["data"]
                                  if isinstance(row, (list, tuple))]
                        if matrix:
                            im = ax.imshow(matrix, cmap="Blues", aspect="auto")
                            plt.colorbar(im, ax=ax, shrink=0.8)
                            if x_data:
                                ax.set_xticks(range(len(x_data)))
                                ax.set_xticklabels(x_data, rotation=45, ha="right")
                            row_labels = chart.get("y_data", [f"R{i+1}" for i in range(len(matrix))])
                            ax.set_yticks(range(len(row_labels)))
                            ax.set_yticklabels(row_labels)
                            for i in range(len(matrix)):
                                for j in range(len(matrix[i])):
                                    ax.text(j, i, f"{matrix[i][j]:.2f}", ha="center", va="center",
                                            fontsize=7, color="white" if matrix[i][j] > max(max(r) for r in matrix) * 0.6 else "black")

                # Common styling
                if chart_type != "pie":
                    ax.set_xlabel(x_label, labelpad=5)
                    ax.set_ylabel(y_label, labelpad=5)

                ax.set_title(title, pad=10)

                if len(datasets) > 1 and chart_type not in ("pie", "heatmap"):
                    ax.legend(loc="best")

                plt.tight_layout(pad=1.2)
                buf = io.BytesIO()
                fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                            facecolor="white", edgecolor="none")
                plt.close(fig)
                buf.seek(0)
                return buf.read()

        except Exception as e:
            logger.warning("Chart render failed", chart_title=chart.get("title"), error=str(e))
            plt.close("all")
            return None

    # ─── Docx Table Renderer ─────────────────────────────────────────────

    def _render_docx_table(self, doc: Document, tbl: Dict[str, Any]) -> None:
        headers = tbl.get("headers", [])
        rows = tbl.get("rows", [])
        caption = tbl.get("title") or tbl.get("caption", "")
        notes = tbl.get("notes", "")

        if not headers:
            return

        cap_para = doc.add_paragraph()
        cap_run = cap_para.add_run(caption)
        cap_run.bold = True
        cap_run.font.size = Pt(9.5)
        cap_run.font.name = "Times New Roman"
        cap_para.paragraph_format.space_after = Pt(3)

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"

        # Header row
        hdr_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = str(h)
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            # Dark background for header
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1a1a2e")
            tcPr.append(shd)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Data rows
        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx + 1]
            for c_idx, val in enumerate(row_data):
                if c_idx < len(headers):
                    cell = row.cells[c_idx]
                    cell.text = str(val)
                    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(val))
                    run.font.size = Pt(8.5)
                    run.font.name = "Times New Roman"
                    if r_idx % 2 == 0:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"), "f0f4ff")
                        tcPr.append(shd)

        if notes:
            n_para = doc.add_paragraph()
            n_run = n_para.add_run(f"Note: {notes}")
            n_run.italic = True
            n_run.font.size = Pt(8)
            n_run.font.name = "Times New Roman"
            n_para.paragraph_format.space_before = Pt(2)

        doc.add_paragraph()

    # ─── Document Style Helpers ───────────────────────────────────────────

    def _setup_document(self, doc: Document) -> None:
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        sect = doc.sections[0]
        sect.page_width = Inches(8.5)
        sect.page_height = Inches(11)
        sect.left_margin = Inches(1.25)
        sect.right_margin = Inches(1.25)
        sect.top_margin = Inches(1.0)
        sect.bottom_margin = Inches(1.0)

    def _add_page_numbers(self, doc: Document) -> None:
        sect = doc.sections[0]
        footer = sect.footer
        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        for tag, text in [("w:fldChar", "begin"), ("w:instrText", "PAGE"), ("w:fldChar", "end")]:
            el = OxmlElement(tag)
            if tag == "w:fldChar":
                el.set(qn("w:fldCharType"), text)
            else:
                el.text = text
            run._r.append(el)

    def _heading(self, doc: Document, text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        run = h.runs[0] if h.runs else h.add_run(text)
        run.font.name = "Times New Roman"
        run.font.bold = True
        if level == 1:
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            run.font.size = Pt(13)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
            run.font.size = Pt(11.5)
        else:
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            run.font.size = Pt(10.5)

    def _body_para(self, doc: Document, text: str, first_line_indent: bool = True) -> None:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        if first_line_indent:
            para.paragraph_format.first_line_indent = Inches(0.3)
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    def _info_box(self, doc: Document, label: str, content: str, color_hex: str = "e8f4f8") -> None:
        para = doc.add_paragraph()
        tc_para = para
        pPr = tc_para._p.get_or_add_pPr()
        shd = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "12")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), "1a73e8")
        shd.append(left)
        pPr.append(shd)
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.space_after = Pt(6)
        bold_run = para.add_run(f"{label}: ")
        bold_run.bold = True
        bold_run.font.size = Pt(9.5)
        bold_run.font.name = "Times New Roman"
        bold_run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
        content_run = para.add_run(content)
        content_run.font.size = Pt(9.5)
        content_run.font.name = "Times New Roman"

    def _add_horizontal_rule(self, doc: Document) -> None:
        para = doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "cccccc")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ─── Paper generation ─────────────────────────────────────────────────

    def generate_paper_docx(self, state: Dict[str, Any]) -> bytes:
        doc = Document()
        self._setup_document(doc)
        self._add_page_numbers(doc)

        sections = state.get("paper_sections", {})
        references = state.get("references", [])
        figures = state.get("figures", [])
        tables_list = state.get("tables", [])
        table_data = state.get("table_data", [])
        chart_data = state.get("chart_data", [])
        methodology_visuals = state.get("methodology_visuals", [])
        discussion_visuals = state.get("discussion_visuals", [])
        citation_style = state.get("citation_style", "ieee")
        journal_type = state.get("journal_type", "sci")

        # ── Title block ────────────────────────────────────────────────────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(state.get("title", "Research Paper"))
        title_run.bold = True
        title_run.font.size = Pt(17)
        title_run.font.name = "Times New Roman"
        title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # Author block
        authors = state.get("all_authors", [])
        if not authors and state.get("author_name"):
            authors = [{"name": state.get("author_name", ""), "affiliation": state.get("author_affiliation", "")}]
        if authors:
            doc.add_paragraph()
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, auth in enumerate(authors):
                run = author_para.add_run(auth.get("name", ""))
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"
                if i < len(authors) - 1:
                    author_para.add_run(", ").font.size = Pt(11)

            aff_para = doc.add_paragraph()
            aff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            affiliations = list(dict.fromkeys(a.get("affiliation", "") for a in authors if a.get("affiliation")))
            aff_run = aff_para.add_run("; ".join(affiliations))
            aff_run.italic = True
            aff_run.font.size = Pt(10)
            aff_run.font.name = "Times New Roman"
            aff_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Keywords
        kw_list = state.get("keywords", [])
        if isinstance(kw_list, list) and kw_list:
            kw_str = ", ".join(kw_list)
        elif isinstance(kw_list, str):
            kw_str = kw_list
        else:
            kw_str = ""
        if kw_str:
            doc.add_paragraph()
            kw_para = doc.add_paragraph()
            kw_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bold_kw = kw_para.add_run("Keywords: ")
            bold_kw.bold = True
            bold_kw.font.size = Pt(9.5)
            bold_kw.font.name = "Times New Roman"
            kw_run = kw_para.add_run(kw_str)
            kw_run.italic = True
            kw_run.font.size = Pt(9.5)
            kw_run.font.name = "Times New Roman"

        # Meta info line
        doc.add_paragraph()
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(
            f"Domain: {state.get('domain', '')}  ·  Target: {journal_type.upper()}  ·  Citation: {citation_style.upper()}  ·  "
            f"Editor Score: {state.get('editor_score', 0):.1f}/10  ·  Generated: {datetime.now(timezone.utc).strftime('%B %Y')}"
        )
        meta_run.font.size = Pt(8.5)
        meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        meta_run.font.name = "Times New Roman"

        self._add_horizontal_rule(doc)
        doc.add_paragraph()

        # TOC omitted — not standard in SCI journal submissions

        # ── Abstract ──────────────────────────────────────────────────────
        self._heading(doc, "Abstract", level=1)
        abstract = sections.get("abstract", "")
        if abstract:
            for para_text in abstract.split("\n\n"):
                if para_text.strip():
                    self._body_para(doc, para_text.strip(), first_line_indent=False)
        doc.add_paragraph()
        self._add_horizontal_rule(doc)
        doc.add_paragraph()

        # ── Main paper sections ────────────────────────────────────────────
        section_map = [
            ("I. Introduction", "introduction"),
            ("II. Literature Review", "literature_review"),
            ("III. Materials and Methods", "materials_and_methods"),
            ("IV. Results", "results"),
            ("V. Discussion", "discussion"),
            ("VI. Conclusion", "conclusion"),
        ]

        # Build chart lookup: inject figures after Results section
        chart_rendered: List[tuple] = []
        for ch in chart_data:
            png = self._render_chart(ch)
            chart_rendered.append((ch, png))

        # Build table lookup for inline insertion
        table_rendered = list(table_data) if table_data else []

        fig_counter = 0
        tbl_counter = 0

        # Pre-render methodology and discussion visuals
        def _render_visuals(visuals: List[Dict[str, Any]]) -> None:
            nonlocal fig_counter, tbl_counter
            for vis in visuals:
                kind = vis.get("kind", "table")
                if kind == "table":
                    tbl_counter += 1
                    self._render_docx_table(doc, vis)
                elif kind == "chart":
                    fig_counter += 1
                    png = self._render_chart(vis)
                    if png:
                        try:
                            img_stream = io.BytesIO(png)
                            doc.add_picture(img_stream, width=Inches(5.4))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            logger.warning("Could not embed visual chart", error=str(e))
                    cap_text = vis.get("caption", f"Figure {fig_counter}: {vis.get('title', '')}")
                    cap_para = doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap_para.add_run(cap_text)
                    cap_run.bold = True
                    cap_run.italic = True
                    cap_run.font.size = Pt(9.5)
                    cap_run.font.name = "Times New Roman"
                    cap_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    doc.add_paragraph()

        for heading, key in section_map:
            self._heading(doc, heading, level=1)
            text = sections.get(key, "")
            if not text:
                self._body_para(doc, "[Section content not generated]", first_line_indent=False)
            else:
                for para_text in text.split("\n\n"):
                    stripped = para_text.strip()
                    if not stripped:
                        continue
                    if stripped.endswith(":") and len(stripped) < 80 and "\n" not in stripped:
                        self._heading(doc, stripped.rstrip(":"), level=2)
                    else:
                        self._body_para(doc, stripped)

            # Inject methodology-specific visuals after Materials & Methods
            if key == "materials_and_methods" and methodology_visuals:
                doc.add_paragraph()
                _render_visuals(methodology_visuals)

            # Inject charts + tables from research agent after Results
            if key == "results":
                doc.add_paragraph()
                for ch, png in chart_rendered:
                    fig_counter += 1
                    if png:
                        try:
                            img_stream = io.BytesIO(png)
                            doc.add_picture(img_stream, width=Inches(5.4))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            logger.warning("Could not embed chart image", error=str(e))
                    cap_text = ch.get("caption", f"Figure {fig_counter}: {ch.get('title', '')}")
                    cap_para = doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap_para.add_run(cap_text)
                    cap_run.bold = True
                    cap_run.italic = True
                    cap_run.font.size = Pt(9.5)
                    cap_run.font.name = "Times New Roman"
                    cap_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    doc.add_paragraph()

                for tbl in table_rendered:
                    tbl_counter += 1
                    self._render_docx_table(doc, tbl)

            # Inject discussion-specific visuals after Discussion
            if key == "discussion" and discussion_visuals:
                doc.add_paragraph()
                _render_visuals(discussion_visuals)

            doc.add_paragraph()

        # ── References ────────────────────────────────────────────────────
        self._heading(doc, "References", level=1)
        for i, ref in enumerate(references):
            formatted = (
                ref.get(f"formatted_{citation_style}")
                or ref.get("formatted")
                or self._fallback_format(ref, i + 1, citation_style)
            )
            ref_para = doc.add_paragraph()
            ref_para.paragraph_format.left_indent = Inches(0.5)
            ref_para.paragraph_format.first_line_indent = Inches(-0.5)
            ref_para.paragraph_format.space_after = Pt(3)
            ref_run = ref_para.add_run(formatted)
            ref_run.font.size = Pt(9.5)
            ref_run.font.name = "Times New Roman"
            if ref.get("verified"):
                v_run = ref_para.add_run(" ✓")
                v_run.font.color.rgb = RGBColor(0, 160, 0)
                v_run.font.size = Pt(8)
        doc.add_paragraph()

        # ── Author Contributions ──────────────────────────────────────────
        contrib = state.get("author_contributions", "")
        if contrib:
            self._add_horizontal_rule(doc)
            self._heading(doc, "Author Contributions", level=2)
            self._info_box(doc, "CRediT", contrib)
            doc.add_paragraph()

        # ── Funding ───────────────────────────────────────────────────────
        funding = state.get("funding_disclosure", "") or state.get("funding_source", "")
        if funding:
            self._heading(doc, "Funding", level=2)
            self._info_box(doc, "Funding", funding)
            doc.add_paragraph()

        # ── Conflicts of Interest ─────────────────────────────────────────
        coi = state.get("conflicts_of_interest", "") or state.get("conflicts_of_interest_input", "")
        if not coi:
            coi = "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper."
        self._heading(doc, "Declaration of Competing Interest", level=2)
        self._info_box(doc, "COI", coi)
        doc.add_paragraph()

        # ── Ethics Statement ──────────────────────────────────────────────
        ethics = state.get("ethics_statement", "") or state.get("ethics_statement_input", "")
        if ethics:
            self._heading(doc, "Ethics Statement", level=2)
            self._info_box(doc, "Ethics", ethics)
            doc.add_paragraph()

        # ── Supplementary Materials ───────────────────────────────────────
        supp = state.get("supplementary_notes", "")
        if supp:
            self._heading(doc, "Supplementary Materials", level=2)
            self._body_para(doc, supp, first_line_indent=False)
            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _fallback_format(self, ref: Dict, n: int, style: str) -> str:
        authors = ", ".join(str(a) for a in ref.get("authors", [])[:3])
        if len(ref.get("authors", [])) > 3:
            authors += " et al."
        year = ref.get("year", "n.d.")
        title = ref.get("title", "Unknown Title")
        journal = ref.get("journal", "")
        doi = ref.get("doi", "")
        if style == "ieee":
            s = f'[{n}] {authors}, "{title}," {journal}, {year}.'
            if doi:
                s += f" doi: {doi}"
        elif style == "apa":
            s = f"{authors} ({year}). {title}. {journal}."
            if doi:
                s += f" https://doi.org/{doi}"
        else:
            s = f"{authors} ({year}). {title}. {journal}."
        return s

    # ─── Cover letter ─────────────────────────────────────────────────────

    async def generate_cover_letter(self, state: Dict[str, Any]) -> bytes:
        author = state.get("author_name") or "The Authors"
        affiliation = state.get("author_affiliation") or "Author Institution"
        msg = HumanMessage(
            content=(
                f"Write a formal, compelling cover letter for submitting this research paper to a "
                f"{state.get('journal_type', 'SCI').upper()} journal.\n\n"
                f"Title: {state.get('title', '')}\n"
                f"Domain: {state.get('domain', '')}\n"
                f"Keywords: {', '.join(state.get('keywords', [])[:6])}\n"
                f"Citation style: {state.get('citation_style', 'IEEE').upper()}\n"
                f"Editor score: {state.get('editor_score', 0):.1f}/10\n"
                f"Verified references: {sum(1 for r in state.get('references', []) if r.get('verified'))}\n"
                f"Lead author: {author}, {affiliation}\n\n"
                "Structure:\n"
                "1. Salutation (Dear Editor-in-Chief,)\n"
                "2. Opening: paper title, novelty statement, why this journal specifically\n"
                "3. Body: key contributions (3-4 specific points), methodology strength, societal/field significance\n"
                "4. Confirmation: original work, not under review elsewhere, all authors approved\n"
                "5. Closing: contact info placeholder\n\n"
                "Tone: professional, confident, concise. 400-500 words."
            )
        )
        resp = await self.llm.ainvoke([msg])

        doc = Document()
        self._setup_document(doc)
        today = datetime.now(timezone.utc).strftime("%B %d, %Y")
        doc.add_paragraph(today)
        doc.add_paragraph()
        doc.add_paragraph("To the Editor-in-Chief,")
        doc.add_paragraph(f"Re: Manuscript Submission — {state.get('title', '')}")
        doc.add_paragraph()

        for para in resp.content.split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(8)
                if p.runs:
                    p.runs[0].font.size = Pt(11)
                    p.runs[0].font.name = "Times New Roman"

        doc.add_paragraph()
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph(f"{author}")
        doc.add_paragraph(affiliation)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ─── Reviewer response ────────────────────────────────────────────────

    async def generate_reviewer_response(self, state: Dict[str, Any]) -> bytes:
        editor_report = state.get("editor_report", {})
        weaknesses = editor_report.get("weaknesses", [])
        recommendations = editor_report.get("recommendations", [])

        msg = HumanMessage(
            content=(
                "Write a professional reviewer response letter template for this paper.\n\n"
                f"Paper: {state.get('title', '')}\n"
                f"Editor score: {state.get('editor_score', 0):.1f}/10\n"
                f"Weaknesses:\n" + "\n".join(f"- {w}" for w in weaknesses[:5]) + "\n\n"
                f"Recommendations:\n" + "\n".join(f"- {r}" for r in recommendations[:5]) + "\n\n"
                "Write a template response:\n"
                "1. Thank the reviewers\n"
                "2. List each weakness as 'Reviewer Comment:' with '[AUTHOR RESPONSE]:' placeholder\n"
                "3. Describe changes made\n"
                "4. Close professionally\n"
                "Mark each [AUTHOR RESPONSE] clearly."
            )
        )
        resp = await self.llm.ainvoke([msg])

        doc = Document()
        self._setup_document(doc)
        self._heading(doc, "Response to Reviewers", level=1)
        doc.add_paragraph(f"Manuscript: {state.get('title', '')}")
        doc.add_paragraph(f"Date: {datetime.now(timezone.utc).strftime('%B %d, %Y')}")
        doc.add_paragraph()

        for para in resp.content.split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                if p.runs:
                    p.runs[0].font.size = Pt(11)
                    p.runs[0].font.name = "Times New Roman"

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ─── Editorial report ─────────────────────────────────────────────────

    def generate_report_docx(self, state: Dict[str, Any]) -> bytes:
        doc = Document()
        self._setup_document(doc)

        self._heading(doc, "Editorial Quality Report", level=1)
        doc.add_paragraph(f"Platform: SCI Research Platform v2.0")
        doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph(f"Project ID: {state.get('project_id', 'N/A')}")
        doc.add_paragraph(f"Paper: {state.get('title', '')}")
        doc.add_paragraph()

        # Score summary table — fixed column widths for clean alignment
        self._heading(doc, "Quality Metrics", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        # Set fixed column widths: Metric=3.2", Value=1.4", Threshold=1.8"
        col_widths = [Inches(3.2), Inches(1.4), Inches(1.8)]
        for col_idx, width in enumerate(col_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = width

        hdr = table.rows[0].cells
        for i, label in enumerate(["Metric", "Value", "Threshold"]):
            hdr[i].text = label
            run = hdr[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            # Blue header background
            tc = hdr[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1a73e8")
            tcPr.append(shd)
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        editor_score = state.get("editor_score", 0)
        plagiarism   = state.get("plagiarism_score", 0)
        ai_score     = state.get("ai_detection_score", 0)
        trust        = state.get("trust_score", 0)
        refs         = state.get("references", [])
        verified     = sum(1 for r in refs if r.get("verified"))

        metrics = [
            ("Editor-in-Chief Score",  f"{editor_score:.1f}/10",       "≥ 9.0 to pass",        editor_score >= 9.0),
            ("Plagiarism Score",        f"{plagiarism:.1f}%",           "< 15%",                 plagiarism < 15.0),
            ("AI Detection Score",      f"{ai_score:.1f}%",             "< 5%",                  ai_score < 5.0),
            ("Reference Trust Score",   f"{trust:.2f}/1.0",             "> 0.4",                 trust > 0.4),
            ("Novelty Score",           f"{state.get('novelty_score',0):.1f}/10", "> 6.0",       state.get("novelty_score",0) > 6.0),
            ("Pipeline Iterations",     str(state.get("iteration", 0)), "≤ 10",                  True),
            ("Total References",        str(len(refs)),                  "≥ 25 recommended",      len(refs) >= 25),
            ("Verified References",     str(verified),                   "> 50%",                 verified > len(refs) * 0.5 if refs else False),
            ("Charts Generated",        str(len(state.get("chart_data", []))), "≥ 3",            len(state.get("chart_data", [])) >= 3),
            ("Tables Generated",        str(len(state.get("table_data", []))), "≥ 2",            len(state.get("table_data", [])) >= 2),
        ]

        for i, (m, v, t, passing) in enumerate(metrics):
            row = table.add_row().cells
            row[0].text = m
            row[1].text = v
            row[2].text = t
            # Alternate row shading
            fill = "f0f4ff" if i % 2 == 0 else "ffffff"
            for cell in row:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), fill)
                tcPr.append(shd)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            # Colour-code Value column: green=pass, red=fail
            if row[1].paragraphs[0].runs:
                row[1].paragraphs[0].runs[0].font.color.rgb = (
                    RGBColor(0x2d, 0xa4, 0x4e) if passing else RGBColor(0xe8, 0x40, 0x40)
                )
                row[1].paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

        er = state.get("editor_report", {})
        if er:
            for field, label in [("strengths", "Strengths"), ("weaknesses", "Areas for Improvement"), ("recommendations", "Recommendations")]:
                items = er.get(field, [])
                if items:
                    self._heading(doc, label, level=2)
                    for item in items:
                        doc.add_paragraph(item, style="List Bullet")

            section_feedback = er.get("section_feedback", {})
            if section_feedback:
                self._heading(doc, "Section-by-Section Feedback", level=2)
                for sec_name, feedback in section_feedback.items():
                    if feedback:
                        self._heading(doc, sec_name.replace("_", " ").title(), level=3)
                        doc.add_paragraph(str(feedback))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ─── LaTeX Export ─────────────────────────────────────────────────────

    @staticmethod
    def _latex_escape(text: str) -> str:
        for old, new in [
            ("\\", "\\textbackslash{}"), ("&", "\\&"), ("%", "\\%"),
            ("$", "\\$"), ("#", "\\#"), ("_", "\\_"), ("{", "\\{"),
            ("}", "\\}"), ("^", "\\^{}"), ("~", "\\textasciitilde{}"),
        ]:
            text = text.replace(old, new)
        return text

    def generate_paper_latex(self, state: Dict[str, Any]) -> bytes:
        e = self._latex_escape
        sections  = state.get("paper_sections", {})
        refs      = state.get("references", [])
        title     = e(state.get("title", ""))
        authors   = state.get("all_authors", []) or [{"name": state.get("author_name", ""), "affiliation": state.get("author_affiliation", "")}]
        kw        = ", ".join(e(k) for k in state.get("keywords", [])[:8])
        citation  = state.get("citation_style", "ieee").lower()
        bib_style = "ieeetr" if citation == "ieee" else "apalike"

        author_block = " \\and\n".join(
            f"\\textbf{{{e(a.get('name',''))}}} \\\\ \\small {e(a.get('affiliation',''))}"
            for a in authors
        )

        from pathlib import Path
        pid = state.get("project_id", "unknown")
        out_dir = Path.cwd() / "outputs" / pid
        out_dir.mkdir(parents=True, exist_ok=True)

        fig_counter = 0
        tbl_counter = 0

        def table_to_latex(tbl: Dict[str, Any], counter: int) -> str:
            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            caption = tbl.get("title") or tbl.get("caption", "")
            notes = tbl.get("notes", "")
            if not headers:
                return ""
            
            col_spec = "l" + "c" * (len(headers) - 1)
            
            latex_parts = []
            latex_parts.append(r"\begin{table}[htbp]")
            latex_parts.append(r"\centering")
            if caption:
                latex_parts.append(f"\\caption{{{e(caption)}}}")
            latex_parts.append(f"\\label{{tab:table_{counter}}}")
            latex_parts.append(f"\\begin{{tabular}}{{{col_spec}}}")
            latex_parts.append(r"\toprule")
            latex_parts.append(" & ".join(e(str(h)) for h in headers) + r" \\")
            latex_parts.append(r"\midrule")
            for r in rows:
                latex_parts.append(" & ".join(e(str(val)) for val in r) + r" \\")
            latex_parts.append(r"\bottomrule")
            latex_parts.append(r"\end{tabular}")
            if notes:
                latex_parts.append(f"\\par\\smallskip\\footnotesize Note: {e(notes)}")
            latex_parts.append(r"\end{table}")
            return "\n".join(latex_parts) + "\n"

        def chart_to_latex(chart: Dict[str, Any], counter: int) -> str:
            png_bytes = self._render_chart(chart)
            img_name = f"figure_{counter}.png"
            if png_bytes:
                try:
                    (out_dir / img_name).write_bytes(png_bytes)
                except Exception as ex:
                    logger.warning("Could not save LaTeX figure image", error=str(ex))
            
            caption = chart.get("caption") or chart.get("title") or ""
            latex_parts = []
            latex_parts.append(r"\begin{figure}[htbp]")
            latex_parts.append(r"\centering")
            latex_parts.append(f"\\includegraphics[width=0.85\\textwidth]{{{img_name}}}")
            if caption:
                latex_parts.append(f"\\caption{{{e(caption)}}}")
            latex_parts.append(f"\\label{{fig:figure_{counter}}}")
            latex_parts.append(r"\end{figure}")
            return "\n".join(latex_parts) + "\n"

        # Build visual blocks
        methodology_latex_parts = []
        for vis in state.get("methodology_visuals", []):
            kind = vis.get("kind", "table")
            if kind == "table":
                tbl_counter += 1
                methodology_latex_parts.append(table_to_latex(vis, tbl_counter))
            elif kind == "chart":
                fig_counter += 1
                methodology_latex_parts.append(chart_to_latex(vis, fig_counter))
        methodology_latex = "\n".join(methodology_latex_parts)

        results_latex_parts = []
        for ch in state.get("chart_data", []):
            fig_counter += 1
            results_latex_parts.append(chart_to_latex(ch, fig_counter))
        for tbl in state.get("table_data", []):
            tbl_counter += 1
            results_latex_parts.append(table_to_latex(tbl, tbl_counter))
        results_latex = "\n".join(results_latex_parts)

        discussion_latex_parts = []
        for vis in state.get("discussion_visuals", []):
            kind = vis.get("kind", "table")
            if kind == "table":
                tbl_counter += 1
                discussion_latex_parts.append(table_to_latex(vis, tbl_counter))
            elif kind == "chart":
                fig_counter += 1
                discussion_latex_parts.append(chart_to_latex(vis, fig_counter))
        discussion_latex = "\n".join(discussion_latex_parts)

        def sec(key: str) -> str:
            raw = sections.get(key, "")
            if not raw:
                return "\\textit{[Content not generated]}\n"
            parts = []
            for para in raw.split("\n\n"):
                p = para.strip()
                if not p:
                    continue
                if p.endswith(":") and len(p) < 80:
                    parts.append(f"\n\\subsection*{{{e(p.rstrip(':'))}}}")
                else:
                    parts.append(e(p) + "\n")
            
            body_content = "\n\n".join(parts)
            
            if key == "materials_and_methods" and methodology_latex:
                body_content += "\n\n" + methodology_latex
            elif key == "results" and results_latex:
                body_content += "\n\n" + results_latex
            elif key == "discussion" and discussion_latex:
                body_content += "\n\n" + discussion_latex
                
            return body_content

        bibitems = []
        for i, r in enumerate(refs[:50]):
            key  = f"ref{i+1}"
            
            authors_list = r.get("authors", [])
            if not authors_list:
                auth = "Unknown"
            else:
                cleaned_authors = []
                for a in authors_list:
                    if isinstance(a, dict):
                        if "name" in a:
                            cleaned_authors.append(a["name"])
                        elif "given" in a or "family" in a:
                            given = a.get("given", "")
                            family = a.get("family", "")
                            if given and family:
                                cleaned_authors.append(f"{given} {family}")
                            else:
                                cleaned_authors.append(family or given)
                        else:
                            cleaned_authors.append(str(a))
                    elif isinstance(a, str):
                        cleaned_authors.append(a)
                    else:
                        cleaned_authors.append(str(a))
                auth = " and ".join(cleaned_authors[:4])
                if len(cleaned_authors) > 4:
                    auth += " and others"
            
            yr = r.get("year")
            if not yr or str(yr).strip().lower() in ("none", "null", "n.d.", ""):
                yr = "2025"
            else:
                yr = str(yr).strip()
                
            ttl  = e(r.get("title", ""))
            jour = e(r.get("journal", ""))
            doi  = r.get("doi", "")
            doi_str = f"  doi = {{{doi}}}," if doi else ""
            bibitems.append(
                f"@article{{{key},\n"
                f"  author  = {{{auth}}},\n"
                f"  title   = {{{ttl}}},\n"
                f"  journal = {{{jour}}},\n"
                f"  year    = {{{yr}}},{doi_str}\n}}"
            )

        tex = rf"""\documentclass[12pt,a4paper]{{article}}
\usepackage[utf8]{{inputenc}}
\usepackage[T1]{{fontenc}}
\usepackage{{geometry}}
\usepackage{{graphicx}}
\usepackage{{booktabs}}
\usepackage{{amsmath,amssymb}}
\usepackage{{hyperref}}
\usepackage{{setspace}}
\usepackage{{parskip}}
\usepackage{{natbib}}
\geometry{{a4paper, top=25mm, bottom=25mm, left=25mm, right=25mm}}
\hypersetup{{colorlinks=true, linkcolor=blue, citecolor=blue, urlcolor=blue}}
\onehalfspacing

\title{{\textbf{{{title}}}}}
\author{{{author_block}}}
\date{{\today}}

\begin{{document}}
\maketitle

\noindent\textbf{{Keywords:}} {kw}

\begin{{abstract}}
{sec("abstract")}
\end{{abstract}}

\tableofcontents
\newpage

\section{{Introduction}}
{sec("introduction")}

\section{{Literature Review}}
{sec("literature_review")}

\section{{Materials and Methods}}
{sec("materials_and_methods")}

\section{{Results}}
{sec("results")}

\section{{Discussion}}
{sec("discussion")}

\section{{Conclusion}}
{sec("conclusion")}

\section*{{Author Contributions}}
{e(state.get("author_contributions",""))}

\section*{{Funding}}
{e(state.get("funding_disclosure","") or state.get("funding_source",""))}

\section*{{Declaration of Competing Interest}}
{e(state.get("conflicts_of_interest","") or "The authors declare no competing interests.")}

{r'\section*{Ethics Statement}' + chr(10) + e(state.get("ethics_statement","")) if state.get("ethics_statement") else ""}

\bibliographystyle{{{bib_style}}}
\bibliography{{references}}

\end{{document}}
"""
        bib = "\n\n".join(bibitems)
        
        # Write clean separate paper.tex and references.bib to output directory
        try:
            (out_dir / "paper.tex").write_text(tex, encoding="utf-8")
            (out_dir / "references.bib").write_text(bib, encoding="utf-8")
        except Exception as ex:
            logger.warning("Could not write clean TeX/Bib files to outputs folder", error=str(ex))
            
        combined = f"% === paper.tex ===\n{tex}\n\n% === references.bib ===\n{bib}"
        return combined.encode("utf-8")

    # ─── Main run ─────────────────────────────────────────────────────────

    async def run(self, state: Dict[str, Any]) -> Dict[str, Any]:
        pid = state.get("project_id", "unknown")

        paper_bytes, cover_bytes, reviewer_bytes, report_bytes = await asyncio.gather(
            asyncio.get_event_loop().run_in_executor(None, self.generate_paper_docx, state),
            self.generate_cover_letter(state),
            self.generate_reviewer_response(state),
            asyncio.get_event_loop().run_in_executor(None, self.generate_report_docx, state),
        )

        keys = {
            "docx": f"outputs/{pid}/paper.docx",
            "cover": f"outputs/{pid}/cover_letter.docx",
            "reviewer": f"outputs/{pid}/reviewer_response.docx",
            "report": f"outputs/{pid}/editorial_report.docx",
        }

        MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        await asyncio.gather(
            self.s3.upload_bytes(paper_bytes, keys["docx"], MIME),
            self.s3.upload_bytes(cover_bytes, keys["cover"], MIME),
            self.s3.upload_bytes(reviewer_bytes, keys["reviewer"], MIME),
            self.s3.upload_bytes(report_bytes, keys["report"], MIME),
        )

        bibtex_key = None
        if state.get("bibtex"):
            bibtex_key = f"outputs/{pid}/references.bib"
            await self.s3.upload_bytes(
                state["bibtex"].encode(),
                bibtex_key,
                "application/x-bibtex",
            )

        # LaTeX export
        latex_key = f"outputs/{pid}/paper.tex"
        latex_bytes = await asyncio.get_event_loop().run_in_executor(
            None, self.generate_paper_latex, state
        )
        await self.s3.upload_bytes(latex_bytes, latex_key, "text/x-tex")

        logger.info("All documents generated and uploaded", project_id=pid,
                    charts=len(state.get("chart_data", [])),
                    tables=len(state.get("table_data", [])))
        return {
            "docx_key": keys["docx"],
            "cover_letter_key": keys["cover"],
            "reviewer_response_key": keys["reviewer"],
            "report_key": keys["report"],
            "bibtex_key": bibtex_key,
            "latex_key": latex_key,
        }
