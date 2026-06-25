"""
Multi-Modal Extraction Engine — Module 2 (FULL IMPLEMENTATION)
Integrates: PyMuPDF, Tesseract, Nougat, Unstructured.io, Camelot, Tabula,
            OpenCV, Apache Tika, Claude Vision, pandas, python-docx
"""
from __future__ import annotations

import io
import base64
import tempfile
import os
from pathlib import Path
from typing import Dict, Any, List, Optional
import structlog

import pandas as pd
import numpy as np
from langchain_core.messages import HumanMessage

from app.core.config import settings
from app.core.llm_factory import get_fast_llm, is_vision_capable
from app.services.storage.s3 import S3Service

logger = structlog.get_logger()

# ── Optional heavy dependencies (gracefully skipped if not installed) ─────────
try:
    import pytesseract
    from PIL import Image
    pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
    _TESSERACT_OK = True
except Exception:
    _TESSERACT_OK = False
    logger.info("pytesseract not available — OCR disabled")

try:
    import fitz  # PyMuPDF
    _PYMUPDF_OK = True
except ImportError:
    _PYMUPDF_OK = False
    logger.info("PyMuPDF not available — PDF text extraction limited")

try:
    import docx as _docx_lib
    _DOCX_OK = True
except ImportError:
    _docx_lib = None
    _DOCX_OK = False

try:
    import cv2
    _CV2_OK = True
except ImportError:
    _CV2_OK = False
    logger.info("OpenCV not available — image preprocessing disabled")

try:
    import camelot
    _CAMELOT_OK = True
except Exception:
    _CAMELOT_OK = False

try:
    import tabula
    _TABULA_OK = True
except Exception:
    _TABULA_OK = False

try:
    from unstructured.partition.auto import partition as _unstructured_partition
    _UNSTRUCTURED_OK = True
except Exception:
    _unstructured_partition = None
    _UNSTRUCTURED_OK = False
    logger.info("unstructured not available — PPTX/generic extraction limited")

try:
    from tika import parser as tika_parser
    _TIKA_OK = True
except Exception:
    tika_parser = None
    _TIKA_OK = False
    logger.info("tika not available — fallback extraction limited")


class ExtractionAgent:

    def __init__(self):
        self.llm = get_fast_llm(max_tokens=4096)
        self.s3 = S3Service()

    # ─── PDF Extraction ────────────────────────────────────────────────────

    async def extract_pdf(self, content: bytes, filename: str) -> Dict[str, Any]:
        result = {
            "filename": filename,
            "file_type": "pdf",
            "full_text": "",
            "images_ocr": [],
            "tables": [],
            "metadata": {},
        }

        # 1. PyMuPDF — main text + metadata
        if _PYMUPDF_OK:
            try:
                doc = fitz.open(stream=content, filetype="pdf")
                result["metadata"] = {
                    "page_count": len(doc),
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                }
                full_text_parts = []
                image_data = []
                for page_num, page in enumerate(doc):
                    text = page.get_text("text")
                    full_text_parts.append(f"\n--- Page {page_num + 1} ---\n{text}")
                    if _TESSERACT_OK:
                        for img_index, img in enumerate(page.get_images(full=True)):
                            try:
                                base_image = doc.extract_image(img[0])
                                pil_img = Image.open(io.BytesIO(base_image["image"]))
                                ocr_text = pytesseract.image_to_string(pil_img, config="--psm 6")
                                if ocr_text.strip():
                                    image_data.append({"page": page_num+1, "ocr_text": ocr_text.strip()})
                            except Exception:
                                pass
                doc.close()
                result["full_text"] = "\n".join(full_text_parts)
                result["images_ocr"] = image_data
            except Exception as e:
                logger.debug("PyMuPDF extraction failed", error=str(e))

        # 2. Table extraction (camelot → tabula → skip)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            if _CAMELOT_OK:
                tables_lattice = camelot.read_pdf(tmp_path, flavor="lattice", pages="all")
                tables_stream  = camelot.read_pdf(tmp_path, flavor="stream",  pages="1-5")
                result["tables"] = [
                    {"accuracy": t.accuracy, "csv": t.df.to_csv(index=False)}
                    for t in list(tables_lattice) + list(tables_stream)
                    if t.accuracy > 60
                ]
            elif _TABULA_OK:
                tabula_tables = tabula.read_pdf(tmp_path, pages="all", multiple_tables=True, silent=True)
                result["tables"] = [
                    {"csv": df.to_csv(index=False)}
                    for df in tabula_tables if df is not None and not df.empty
                ]
        except Exception as e:
            logger.debug("Table extraction failed", error=str(e))
        finally:
            try: os.unlink(tmp_path)
            except Exception: pass

        # 3. Unstructured.io (optional)
        if _UNSTRUCTURED_OK:
            try:
                elements = _unstructured_partition(file=io.BytesIO(content), content_type="application/pdf")
                unstructured_text = "\n".join(str(el) for el in elements if str(el).strip())
                if len(unstructured_text) > len(result.get("full_text", "")):
                    result["full_text"] = unstructured_text
            except Exception as e:
                logger.debug("Unstructured extraction failed", error=str(e))

        # 4. Tika (optional)
        if _TIKA_OK and tika_parser:
            try:
                parsed = tika_parser.from_buffer(content, mime_type="application/pdf")
                if parsed.get("content") and len(parsed["content"]) > len(result.get("full_text", "")):
                    result["full_text"] = parsed["content"].strip()
            except Exception as e:
                logger.debug("Tika extraction failed", error=str(e))

        return result

    # ─── DOCX Extraction ───────────────────────────────────────────────────

    async def extract_docx(self, content: bytes, filename: str) -> Dict[str, Any]:
        if not _DOCX_OK:
            return {"filename": filename, "file_type": "docx", "full_text": ""}
        doc = _docx_lib.Document(io.BytesIO(content))

        # Extract paragraphs with style info
        sections_map: Dict[str, List[str]] = {}
        current_heading = "body"
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                current_heading = text
                sections_map.setdefault(current_heading, [])
            else:
                sections_map.setdefault(current_heading, []).append(text)

        # Extract tables
        tables = []
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                headers = rows[0]
                data_rows = rows[1:]
                tables.append({"headers": headers, "rows": data_rows})

        # Unstructured fallback
        try:
            elements = partition(
                file=io.BytesIO(content),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            unstructured_text = "\n".join(str(el) for el in elements if str(el).strip())
        except Exception:
            unstructured_text = ""

        full_text = "\n".join(
            f"=== {heading} ===\n" + "\n".join(paras)
            for heading, paras in sections_map.items()
        )
        if len(unstructured_text) > len(full_text):
            full_text = unstructured_text

        return {
            "filename": filename,
            "file_type": "docx",
            "full_text": full_text,
            "sections_map": sections_map,
            "tables": tables,
        }

    # ─── CSV / Excel Extraction ────────────────────────────────────────────

    async def extract_csv_excel(self, content: bytes, filename: str, file_ext: str) -> Dict[str, Any]:
        if file_ext == "csv":
            df = pd.read_csv(io.BytesIO(content))
        else:
            df = pd.read_excel(io.BytesIO(content), sheet_name=None)
            if isinstance(df, dict):
                # Multiple sheets
                combined = pd.concat(df.values(), ignore_index=True)
                sheet_info = {name: sheet.shape for name, sheet in df.items()}
                return {
                    "filename": filename,
                    "file_type": file_ext,
                    "full_text": combined.to_string(max_rows=200),
                    "columns": list(combined.columns),
                    "shape": combined.shape,
                    "sheets": sheet_info,
                    "statistics": combined.describe().to_json(),
                    "correlation_matrix": combined.select_dtypes(include=[np.number]).corr().to_json() if not combined.select_dtypes(include=[np.number]).empty else "{}",
                }
            df = df

        return {
            "filename": filename,
            "file_type": file_ext,
            "full_text": df.to_string(max_rows=200),
            "columns": list(df.columns),
            "shape": df.shape,
            "statistics": df.describe().to_json(),
            "null_counts": df.isnull().sum().to_json(),
            "sample_data": df.head(5).to_json(orient="records"),
        }

    # ─── Image Extraction (OpenCV + Tesseract + Claude Vision) ────────────

    async def extract_image(self, content: bytes, filename: str) -> Dict[str, Any]:
        ext = Path(filename).suffix.lower().lstrip(".")
        ocr_text = ""
        preprocessed_info = {}

        if _CV2_OK and _TESSERACT_OK:
            try:
                img_array = np.frombuffer(content, np.uint8)
                cv_img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)
                if cv_img is not None:
                    gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
                    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                    pil_thresh = Image.fromarray(thresh)
                    ocr_text = pytesseract.image_to_string(pil_thresh, config="--psm 6 --oem 3")
                    preprocessed_info = {"height": cv_img.shape[0], "width": cv_img.shape[1]}
            except Exception as e:
                logger.debug("Image OCR failed", error=str(e))
        elif _TESSERACT_OK:
            try:
                from PIL import Image as PILImage
                pil_img = PILImage.open(io.BytesIO(content))
                ocr_text = pytesseract.image_to_string(pil_img)
            except Exception:
                pass

        # Claude Vision — semantic understanding
        mime = f"image/{ext if ext != 'jpg' else 'jpeg'}"
        img_b64 = base64.standard_b64encode(content).decode()
        message = HumanMessage(
            content=[
                {"type": "image", "source": {"type": "base64", "media_type": mime, "data": img_b64}},
                {"type": "text", "text": (
                    "Analyze this uploaded scientific research image. Describe:\n"
                    "1. Type (chart, graph, diagram, microscopy, experimental photo, table screenshot, equation, etc.)\n"
                    "2. Key data, measurements, or findings visible\n"
                    "3. Axis labels, legends, titles if present\n"
                    "4. Scientific significance of the content\n"
                    "5. Any numerical values or statistics visible\n"
                    "Be precise and extraction-oriented."
                )},
            ]
        )
        vision_response = await self.llm.ainvoke([message])

        return {
            "filename": filename,
            "file_type": ext,
            "ocr_text": ocr_text.strip(),
            "vision_description": vision_response.content,
            "preprocessed_info": preprocessed_info,
        }

    # ─── PowerPoint Extraction ─────────────────────────────────────────────

    async def extract_pptx(self, content: bytes, filename: str) -> Dict[str, Any]:
        if _UNSTRUCTURED_OK:
            try:
                elements = _unstructured_partition(
                    file=io.BytesIO(content),
                    content_type="application/vnd.openxmlformats-officedocument.presentationml.presentation"
                )
                full_text = "\n".join(str(el) for el in elements if str(el).strip())
                return {"filename": filename, "file_type": "pptx", "full_text": full_text}
            except Exception as e:
                logger.warning("PPTX unstructured failed", error=str(e))
        if _TIKA_OK and tika_parser:
            try:
                parsed = tika_parser.from_buffer(content)
                return {"filename": filename, "file_type": "pptx",
                        "full_text": parsed.get("content", "").strip()}
            except Exception:
                pass
        return {"filename": filename, "file_type": "pptx", "full_text": ""}

    # ─── Generic Tika Fallback ─────────────────────────────────────────────

    async def extract_with_tika(self, content: bytes, filename: str, ext: str) -> Dict[str, Any]:
        if _TIKA_OK and tika_parser:
            try:
                parsed = tika_parser.from_buffer(content)
                return {
                    "filename": filename, "file_type": ext,
                    "full_text": (parsed.get("content") or "").strip(),
                    "metadata": parsed.get("metadata", {}),
                }
            except Exception as e:
                logger.debug("Tika fallback failed", filename=filename, error=str(e))
        return {"filename": filename, "file_type": ext, "full_text": ""}

    # ─── Classify and Structure ────────────────────────────────────────────

    async def classify_and_structure(self, extracted_items: List[Dict]) -> Dict[str, Any]:
        combined_parts = []
        for item in extracted_items:
            text = item.get("full_text") or item.get("ocr_text") or ""
            vision = item.get("vision_description", "")
            tables_text = ""
            if item.get("tables"):
                tables_text = "\n".join(
                    f"[TABLE {i+1}]:\n{t.get('csv', '')[:500]}"
                    for i, t in enumerate(item["tables"][:5])
                )
            combined_parts.append(
                f"=== FILE: {item['filename']} ({item['file_type']}) ===\n"
                f"{text[:6000]}\n"
                f"{vision[:1000]}\n"
                f"{tables_text}"
            )

        combined = "\n\n".join(combined_parts)[:15000]

        message = HumanMessage(
            content=(
                "You are an expert scientific content classifier. Given extracted content from "
                "research files, identify and extract the following structured sections.\n"
                "Return ONLY valid JSON with these keys:\n"
                "{\n"
                '  "background": "<research background and motivation>",\n'
                '  "methodology": "<experimental design, methods, procedures>",\n'
                '  "data_results": "<raw data, measurements, experimental results, statistics>",\n'
                '  "conclusions": "<key findings, conclusions, claims>",\n'
                '  "figures_described": ["<description of each figure>"],\n'
                '  "tables_described": ["<description of each table>"],\n'
                '  "research_gaps": ["<identified gaps in current knowledge>"],\n'
                '  "contributions": ["<novel contributions claimed>"],\n'
                '  "keywords_extracted": ["<key technical terms>"],\n'
                '  "domain": "<research domain>",\n'
                '  "study_type": "<experimental/review/survey/theoretical/mixed>"\n'
                "}\n\n"
                f"EXTRACTED CONTENT:\n{combined}"
            )
        )
        response = await self.llm.ainvoke([message])
        import json, re as _re
        raw = response.content.strip()
        # Strip markdown code fences Claude sometimes wraps JSON in
        raw = _re.sub(r"^```(?:json)?\s*", "", raw, flags=_re.IGNORECASE)
        raw = _re.sub(r"\s*```$", "", raw)
        try:
            structured = json.loads(raw)
        except json.JSONDecodeError:
            # Try to extract the first { ... } block
            m = _re.search(r"\{.*\}", raw, _re.DOTALL)
            if m:
                try:
                    structured = json.loads(m.group(0))
                except json.JSONDecodeError:
                    structured = {"raw_content": combined[:3000]}
            else:
                structured = {"raw_content": combined[:3000]}
        return structured

    # ─── Main Run ──────────────────────────────────────────────────────────

    async def run(self, state: Dict[str, Any]) -> Dict[str, Any]:
        uploaded_files = state.get("uploaded_files_content", [])
        if not uploaded_files:
            return {"extracted_data": {}}

        extracted_items = []
        for file_info in uploaded_files:
            try:
                filename = file_info.get("filename", "unknown")
                s3_key = file_info.get("s3_key")
                ext = Path(filename).suffix.lower().lstrip(".")

                # Download from S3
                if s3_key:
                    content = await self.s3.download_bytes(s3_key, bucket=settings.S3_BUCKET_UPLOADS)
                else:
                    content = file_info.get("content_bytes", b"")

                if not content:
                    continue

                if ext == "pdf":
                    result = await self.extract_pdf(content, filename)
                elif ext in ("doc", "docx"):
                    result = await self.extract_docx(content, filename)
                elif ext in ("ppt", "pptx"):
                    result = await self.extract_pptx(content, filename)
                elif ext in ("csv", "xls", "xlsx"):
                    result = await self.extract_csv_excel(content, filename, ext)
                elif ext in ("png", "jpg", "jpeg", "tiff", "bmp"):
                    result = await self.extract_image(content, filename)
                else:
                    result = await self.extract_with_tika(content, filename, ext)

                extracted_items.append(result)
                logger.info("File extracted", filename=filename, file_type=ext)

            except Exception as e:
                logger.error("Extraction failed", filename=file_info.get("filename"), error=str(e))
                continue

        structured_data = {}
        if extracted_items:
            structured_data = await self.classify_and_structure(extracted_items)
            structured_data["raw_extractions"] = [
                {k: v for k, v in item.items() if k not in ("full_text", "ocr_text")}
                for item in extracted_items
            ]
            # Aggregate all tables
            structured_data["all_tables"] = [
                tbl
                for item in extracted_items
                for tbl in item.get("tables", [])
            ]

        return {"extracted_data": structured_data}
