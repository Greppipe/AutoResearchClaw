"""
Local Pipeline Runner  —  checkpoint/resume + live preview + cost estimation
=============================================================================
Runs the full SCI Research Platform pipeline on localhost without requiring
PostgreSQL, Redis, MinIO/S3, Celery, or Clerk.

Key features:
  • S3 replaced with local filesystem writes
  • SqliteSaver checkpoint: crash → resume from last completed step
  • Enhanced progress callback: (message, pct, preview_content)
  • Cost estimator before any API calls are made
"""
from __future__ import annotations

import asyncio
import json
import os
import sys
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

# ── 1. Bootstrap required env-vars BEFORE any app.* imports ─────────────────
_DUMMY: Dict[str, str] = {
    "SECRET_KEY":            "local-dev-key-placeholder-32-chars!!",
    "DATABASE_URL":          "postgresql://local:local@localhost:5432/local",
    "S3_ACCESS_KEY":         "local",
    "S3_SECRET_KEY":         "local",
    "CLERK_SECRET_KEY":      "sk_local_placeholder",
    "CLERK_PUBLISHABLE_KEY": "pk_local_placeholder",
    "CLERK_JWKS_URL":        "http://localhost",
    # Always set a dummy Anthropic key so pydantic-settings validation passes
    # even when using Groq / Gemini / Ollama
    "ANTHROPIC_API_KEY":     "sk-ant-dummy-not-used",
}
for _k, _v in _DUMMY.items():
    os.environ.setdefault(_k, _v)
os.environ.setdefault("TAVILY_API_KEY", "tvly-local-placeholder")

# Windows Tesseract auto-detect
if sys.platform == "win32" and not os.environ.get("TESSERACT_CMD"):
    _win_tess = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if os.path.exists(_win_tess):
        os.environ["TESSERACT_CMD"] = _win_tess

# Stub aioboto3 so s3.py can be imported without the real package installed.
# _patch_s3() replaces S3Service entirely before any agent uses it, so the
# stub's Session is never actually called in local mode.
import types as _types
if "aioboto3" not in sys.modules:
    _aioboto3_stub = _types.ModuleType("aioboto3")
    class _StubSession:
        def client(self, *a, **kw): return self
        async def __aenter__(self): return self
        async def __aexit__(self, *a): pass
    _aioboto3_stub.Session = _StubSession  # type: ignore[attr-defined]
    sys.modules["aioboto3"] = _aioboto3_stub

if "botocore" not in sys.modules:
    _botocore_stub = _types.ModuleType("botocore")
    _botocore_config_stub = _types.ModuleType("botocore.config")
    class _StubConfig:
        def __init__(self, *a, **kw): pass
    _botocore_config_stub.Config = _StubConfig
    _botocore_stub.config = _botocore_config_stub
    sys.modules["botocore"] = _botocore_stub
    sys.modules["botocore.config"] = _botocore_config_stub

import structlog
logger = structlog.get_logger()

# ── 2. Callback type ──────────────────────────────────────────────────────────
# progress_callback(message, percent, preview_content)
ProgressCB = Callable[[str, float, str], None]

# ── 3. Local filesystem S3 replacement ──────────────────────────────────────

class _LocalS3:
    def __init__(self, base_dir: Path):
        self._base = base_dir

    def _path(self, key: str) -> Path:
        p = self._base / key
        p.parent.mkdir(parents=True, exist_ok=True)
        return p

    async def upload_bytes(self, data: bytes, key: str,
                           content_type: str, bucket: Optional[str] = None) -> str:
        self._path(key).write_bytes(data)
        return key

    async def upload_file(self, file_path: str, key: str,
                          content_type: str, bucket: Optional[str] = None) -> str:
        import shutil
        shutil.copy2(file_path, self._path(key))
        return key

    async def download_bytes(self, key: str, bucket: Optional[str] = None) -> bytes:
        p = self._path(key)
        if p.exists():
            return p.read_bytes()
        raise FileNotFoundError(f"LocalS3: {p}")

    async def generate_presigned_url(self, key: str, **_) -> str:
        return f"file://{self._path(key)}"

    async def delete_object(self, key: str, **_) -> None:
        p = self._path(key)
        if p.exists():
            p.unlink()

    async def ensure_buckets(self) -> None:
        pass


def _patch_s3(out_dir: Path) -> None:
    import importlib
    cls = type("LocalS3Service", (_LocalS3,),
                {"__init__": lambda self: _LocalS3.__init__(self, out_dir)})
    for mod_path in ("app.services.storage.s3",
                     "app.agents.document.agent",
                     "app.agents.extraction.agent"):
        try:
            mod = importlib.import_module(mod_path)
            mod.S3Service = cls  # type: ignore
        except ImportError:
            pass


# ── 4. Cost estimator ─────────────────────────────────────────────────────────

def estimate_cost(params: Dict[str, Any], file_count: int) -> Dict[str, Any]:
    """
    Rough token + cost estimate before the run starts.
    Opus 4.7:  $15 / 1M input  •  $75 / 1M output
    Sonnet 4.6: $3 / 1M input  •  $15 / 1M output
    """
    wc           = params.get("preferred_word_count", 8000)
    context_in   = 4_000 + file_count * 5_000        # system prompts + user context
    paper_out    = int(wc * 1.35)                     # ~1.35 tokens per word
    agents_in    = 20_000                             # authentication, audit, editor, etc.
    agents_out   = 6_000                              # structured JSON responses
    humanize_in  = int(paper_out * 0.8)              # re-reads the draft
    humanize_out = int(paper_out * 0.6)              # rewrites flagged sections

    total_in  = context_in + agents_in + humanize_in
    total_out = paper_out  + agents_out + humanize_out

    # Opus 4.7 handles writing; Sonnet 4.6 handles fast tasks (roughly 60/40 split)
    opus_in   = total_in  * 0.6;  sonnet_in  = total_in  * 0.4
    opus_out  = total_out * 0.6;  sonnet_out = total_out * 0.4

    cost = (opus_in  / 1e6 * 15  + opus_out  / 1e6 * 75 +
            sonnet_in / 1e6 * 3  + sonnet_out / 1e6 * 15)

    return {
        "input_tokens":        int(total_in),
        "output_tokens":       int(total_out),
        "estimated_usd":       round(cost, 2),
        "low_usd":             round(cost * 0.65, 2),
        "high_usd":            round(cost * 1.9,  2),
    }


# ── 5. State builder ──────────────────────────────────────────────────────────

def _build_state(project_id: str, params: Dict[str, Any],
                 files: List[Dict[str, Any]]) -> Dict[str, Any]:
    return {
        "project_id": project_id, "user_id": "local",
        "title":               params.get("title", ""),
        "domain":              params.get("domain", ""),
        "keywords":            params.get("keywords", []),
        "objectives":          params.get("objectives", ""),
        "problem_statement":   params.get("problem_statement", ""),
        "research_gap":        params.get("research_gap", ""),
        "hypothesis":          params.get("hypothesis", ""),
        "novel_contribution":  params.get("novel_contribution", ""),
        "scope":               params.get("scope", ""),
        "study_type":          params.get("study_type", ""),
        "journal_type":        params.get("journal_type", "SCI"),
        "citation_style":      params.get("citation_style", "ieee"),
        "preferred_word_count":int(params.get("preferred_word_count", 8000)),
        "writing_tone":        params.get("writing_tone", "academic"),
        "additional_instructions": params.get("additional_instructions", ""),
        "author_name":         params.get("author_name", ""),
        "author_affiliation":  params.get("author_affiliation", ""),
        "all_authors":         params.get("all_authors", []),
        "funding_source":      params.get("funding_source", ""),
        "conflicts_of_interest_input": params.get("conflicts_of_interest", ""),
        "ethics_statement_input":      params.get("ethics_statement", ""),
        "methodology_description":     params.get("methodology_description", ""),
        "dataset_description":         params.get("dataset_description", ""),
        "analysis_methods":            params.get("analysis_methods", ""),
        "tools_used":                  params.get("tools_used", ""),
        "expected_findings":           params.get("expected_findings", ""),
        "research_significance":       params.get("research_significance", ""),
        "uploaded_files_content": files,
        "extracted_data": {}, "paper_sections": {},
        "suggested_figures": [], "suggested_tables": [],
        "references": [], "figures": [], "tables": [], "toc": [], "bibtex": None,
        "materials_and_methods": "", "author_contributions": "",
        "funding_disclosure": "", "conflicts_of_interest": "",
        "ethics_statement": "", "supplementary_notes": "",
        "chart_data": [], "table_data": [],
        "methodology_visuals": [], "discussion_visuals": [],
        "ddi_brief": {}, "ddi_chunks": 0, "ddi_grounded": False,
        # Module 0: populated by JournalIntelligenceAgent before any drafting
        "journal_constraints": {},
        # Module 1.5: populated by EvidenceMapperAgent before ResearchAgent writes
        "evidence_anchors": [],
        "plagiarism_score": 0.0, "ai_detection_score": 0.0,
        "trust_score": 0.0, "novelty_score": 0.0,
        "editor_score": 0.0, "editor_report": {},
        "modules_to_retry": [],
        # Structured repair instructions from EditorAgent — injected into agents on retry
        "repair_jobs": [],
        "iteration": 0,
        "current_step": "init", "pipeline_complete": False, "error": None,
        "output_docx_key": None, "output_cover_letter_key": None,
        "output_reviewer_key": None, "output_report_key": None,
        "output_bibtex_key": None,
        "progress_log": [], "messages": [],
    }


# ── 6. Checkpoint helpers ─────────────────────────────────────────────────────

def _get_checkpointer(db_path: Path):
    """Return a SqliteSaver if available, else None (no checkpointing)."""
    try:
        from langgraph.checkpoint.sqlite import SqliteSaver
        return SqliteSaver.from_conn_string(str(db_path))
    except Exception:
        try:
            from langgraph.checkpoint.memory import MemorySaver
            return MemorySaver()
        except Exception:
            return None


def scan_incomplete_runs(outputs_root: Optional[Path] = None) -> List[Dict[str, Any]]:
    """
    Scan ./outputs/ for pipeline runs that started but never set pipeline_complete=True.
    Returns list of {project_id, title, started_at, last_step, db_path}.
    """
    root = outputs_root or (Path.cwd() / "outputs")
    incomplete = []
    if not root.exists():
        return incomplete

    for db_file in root.rglob("checkpoint.db"):
        try:
            from langgraph.checkpoint.sqlite import SqliteSaver
            cp = SqliteSaver.from_conn_string(str(db_file))
            # Look for any thread that isn't complete
            pid = db_file.parent.name
            cfg = {"configurable": {"thread_id": pid}}
            saved = cp.get(cfg)
            if saved:
                cv = saved.get("channel_values", {})
                if not cv.get("pipeline_complete", False):
                    incomplete.append({
                        "project_id":  pid,
                        "title":       cv.get("title", "Unknown"),
                        "last_step":   cv.get("current_step", "unknown"),
                        "started_at":  db_file.stat().st_ctime,
                        "db_path":     str(db_file),
                        "out_dir":     str(db_file.parent),
                    })
        except Exception:
            continue
    return incomplete


# ── 7. Full paper pipeline ───────────────────────────────────────────────────

async def run_paper(
    params: Dict[str, Any],
    file_paths: List[str],
    output_dir: Optional[Path] = None,
    progress_callback: Optional[ProgressCB] = None,
    resume_project_id: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Run the full SCI paper pipeline locally with checkpoint/resume support.

    Args:
        params:             research parameters from CLI wizard
        file_paths:         local paths to uploaded files
        output_dir:         where to write outputs (default: ./outputs/<id>)
        progress_callback:  fn(message, percent, preview_content)
        resume_project_id:  project_id of an incomplete run to resume

    Returns:
        dict with project_id, output_dir, scores, file paths
    """
    project_id = resume_project_id or str(uuid.uuid4())[:8]
    out_dir    = output_dir or (Path.cwd() / "outputs" / project_id)
    out_dir.mkdir(parents=True, exist_ok=True)

    _patch_s3(out_dir)

    file_contents: List[Dict[str, Any]] = []
    for fp in file_paths:
        p = Path(fp)
        if p.exists():
            file_contents.append({"filename": p.name, "content_bytes": p.read_bytes()})
        else:
            logger.warning("File not found", path=fp)

    # Build checkpointer
    checkpointer = _get_checkpointer(out_dir / "checkpoint.db")

    # Import and compile pipeline
    from app.agents.orchestrator.graph import build_pipeline
    if checkpointer:
        local_graph = build_pipeline().compile(checkpointer=checkpointer)
    else:
        local_graph = build_pipeline().compile()

    config = {"configurable": {"thread_id": project_id}} if checkpointer else {}

    # Detect resume
    resuming = False
    if checkpointer and resume_project_id:
        try:
            saved = checkpointer.get(config)
            if saved and saved.get("channel_values", {}).get("current_step"):
                resuming = True
        except Exception:
            pass

    initial_state = None if resuming else _build_state(project_id, params, file_contents)

    def _cb(msg: str, pct: float, content: str = ""):
        if progress_callback:
            progress_callback(msg, pct, content)

    _cb("Initialising pipeline...", 1.0)

    last_pct    = 0.0
    final_state = initial_state or {}
    preview     = ""

    async for chunk in local_graph.astream(initial_state, config=config):
        for node_name, node_state in chunk.items():
            if not isinstance(node_state, dict):
                continue
            final_state = {**final_state, **node_state}

            # Build preview content from latest completed section
            sections = final_state.get("paper_sections", {})
            for key in ("abstract", "introduction", "results", "discussion"):
                if sections.get(key):
                    preview = sections[key][:600].strip()
                    break

            logs = node_state.get("progress_log", [])
            if logs:
                last_log = logs[-1]
                pct  = float(last_log.get("progress_percent", last_pct))
                msg  = last_log.get("message", node_name)
                if pct >= last_pct:
                    _cb(msg, pct, preview)
                    last_pct = pct

    _cb("Pipeline complete!", 100.0, preview)

    # Map output keys to local paths
    def _resolve(key_field: str, filename: str) -> Optional[str]:
        key = final_state.get(key_field)
        if not key:
            return None
        direct = out_dir / Path(key).name
        if direct.exists():
            return str(direct)
        via_key = Path.cwd() / "outputs" / key
        if via_key.exists():
            return str(via_key)
        return None

    return {
        "project_id":        project_id,
        "output_dir":        str(out_dir),
        "paper":             _resolve("output_docx_key",        "paper.docx"),
        "cover letter":      _resolve("output_cover_letter_key","cover_letter.docx"),
        "editorial report":  _resolve("output_report_key",      "editorial_report.docx"),
        "latex source":      str(out_dir / "paper.tex") if (out_dir / "paper.tex").exists() else None,
        "editor_score":      final_state.get("editor_score", 0.0),
        "plagiarism_score":  final_state.get("plagiarism_score", 0.0),
        "ai_detection_score":final_state.get("ai_detection_score", 0.0),
        "trust_score":       final_state.get("trust_score", 0.0),
        "novelty_score":     final_state.get("novelty_score", 0.0),
        "reference_count":   len(final_state.get("references", [])),
        "chart_count":       len(final_state.get("chart_data", [])),
        "table_count":       len(final_state.get("table_data", [])),
    }


# ── 8. Summary pipeline ──────────────────────────────────────────────────────

async def run_summary(
    params: Dict[str, Any],
    file_paths: List[str],
    output_dir: Optional[Path] = None,
    progress_callback: Optional[ProgressCB] = None,
) -> Dict[str, Any]:
    """
    Run extract → DDI → SummaryAgent locally. Faster than full paper (~3-8 min).
    """
    project_id = str(uuid.uuid4())[:8]
    out_dir    = output_dir or (Path.cwd() / "outputs" / project_id)
    out_dir.mkdir(parents=True, exist_ok=True)
    _patch_s3(out_dir)

    file_contents = [
        {"filename": Path(fp).name, "content_bytes": Path(fp).read_bytes()}
        for fp in file_paths if Path(fp).exists()
    ]
    state = _build_state(project_id, params, file_contents)

    def _cb(msg: str, pct: float, content: str = ""):
        if progress_callback:
            progress_callback(msg, pct, content)

    _cb("Extracting documents...", 10.0)
    from app.agents.extraction.agent import ExtractionAgent
    state["extracted_data"] = (await ExtractionAgent().run(state)).get("extracted_data", {})

    _cb("Deep Document Intelligence — embedding & interrogating...", 30.0)
    from app.agents.notebook.agent import DeepDocumentIntelligence
    state.update(await DeepDocumentIntelligence().interrogate(state))

    _cb("Generating research summary...", 60.0)
    from app.agents.summary.agent import SummaryAgent
    summary = await SummaryAgent().run(state, state.get("ddi_brief", {}))

    _cb("Writing summary document...", 85.0)
    summary_path = _write_summary_docx(summary, state, out_dir)

    _cb("Summary complete!", 100.0, summary.get("one_sentence_summary", ""))

    return {
        "project_id":  project_id,
        "output_dir":  str(out_dir),
        "summary_docx":str(summary_path),
        **summary,
    }


# ── 9. Summary Word document renderer ────────────────────────────────────────

def _write_summary_docx(summary: Dict[str, Any],
                        state:   Dict[str, Any],
                        out_dir: Path) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime, timezone

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1.0)
        sec.left_margin = sec.right_margin = Inches(1.25)

    def h(text, level=1):
        p = doc.add_heading(text, level=level)
        for r in p.runs:
            r.font.name = "Times New Roman"

    def body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size, r.font.name = Pt(11), "Times New Roman"

    def bullet(text):
        p = doc.add_paragraph(text, style="List Bullet")
        if p.runs:
            p.runs[0].font.size, p.runs[0].font.name = Pt(11), "Times New Roman"

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Research Summary Report")
    r.bold, r.font.size, r.font.name = True, Pt(18), "Times New Roman"
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    paper_title = state.get("title", "")
    if paper_title:
        doc.add_paragraph()
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = s.add_run(paper_title)
        sr.bold, sr.font.size, sr.font.name = True, Pt(13), "Times New Roman"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Domain: {state.get('domain','')}  ·  "
        f"Generated: {datetime.now(timezone.utc).strftime('%B %Y')}"
    )
    mr.font.size, mr.font.color.rgb, mr.font.name = (
        Pt(9), RGBColor(0x88, 0x88, 0x88), "Times New Roman"
    )
    doc.add_paragraph()

    # One-sentence summary callout
    one = summary.get("one_sentence_summary", "")
    if one:
        box = doc.add_paragraph()
        box.paragraph_format.left_indent = Inches(0.3)
        box.paragraph_format.right_indent = Inches(0.3)
        pPr = box._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:color"), "1a73e8")
        pBdr.append(left)
        pPr.append(pBdr)
        r = box.add_run(one)
        r.italic, r.font.size, r.font.name = True, Pt(11), "Times New Roman"
    doc.add_paragraph()

    for label, key in [
        ("Executive Summary",         "executive_summary"),
        ("Research Background",       "research_background"),
        ("Methodology Overview",      "methodology_overview"),
        ("Research Gap & Contribution","research_gap_and_contribution"),
        ("Impact Statement",          "impact_statement"),
    ]:
        val = summary.get(key, "")
        if val:
            h(label)
            body(val)
            doc.add_paragraph()

    if summary.get("key_findings"):
        h("Key Findings")
        for f in summary["key_findings"]:
            bullet(f)
        doc.add_paragraph()

    if summary.get("key_statistics"):
        h("Key Statistics")
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, lbl in enumerate(["Metric", "Value", "Context"]):
            hdr[i].text = lbl
            if hdr[i].paragraphs[0].runs:
                hdr[i].paragraphs[0].runs[0].bold = True
        for s in summary["key_statistics"]:
            row = tbl.add_row().cells
            row[0].text, row[1].text, row[2].text = (
                str(s.get("metric","")), str(s.get("value","")), str(s.get("context",""))
            )
        doc.add_paragraph()

    if summary.get("limitations"):
        h("Limitations")
        for lim in summary["limitations"]:
            bullet(lim)
        doc.add_paragraph()

    if summary.get("future_work"):
        h("Future Work")
        for fw in summary["future_work"]:
            bullet(fw)
        doc.add_paragraph()

    if summary.get("glossary"):
        h("Glossary")
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, lbl in enumerate(["Term", "Definition"]):
            hdr[i].text = lbl
            if hdr[i].paragraphs[0].runs:
                hdr[i].paragraphs[0].runs[0].bold = True
        for g in summary["glossary"]:
            row = tbl.add_row().cells
            row[0].text, row[1].text = str(g.get("term","")), str(g.get("definition",""))
        doc.add_paragraph()

    out = out_dir / "research_summary.docx"
    doc.save(str(out))
    return out
